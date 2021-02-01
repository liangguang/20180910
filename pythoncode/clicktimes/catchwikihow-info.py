# -*- coding: utf-8 -*-

from docx import Document
import json
import re
import os
import sys
import shutil
import time
import traceback
import random
import hashlib
from contextlib import closing
from urllib.request import urlopen
from bs4 import BeautifulSoup

def parseFile(file):
	content = ''
	with open(file,encoding='utf-8') as file_obj:
		content = file_obj.read()
	#print(content)
	soup = BeautifulSoup(content, 'lxml')
	#创建word
	Doc = Document()
	#标题
	title = soup.find('h1',id='section_0').text
	Doc.add_heading(title)
	#提要
	tiyao =  soup.find('div',id='mf-section-0').text
	Doc.add_heading('提要')
	Doc.add_paragraph(tiyao)
	#开始段落
	newslist = soup.findAll('div',class_='step')
	#print(newslist)
	news_url = []
	for news in newslist:
		#print(news)
		bt = news.find('b').text
		#排除某些标签
		[s.extract() for s in news('sub')] 
		dl = news.text
		Doc.add_heading(bt)
		Doc.add_paragraph(dl)
	Doc.save('D:\ps\words\\' + title+'.docx')

if __name__ == '__main__':
	rootdir = 'D:\ps'
	list = os.listdir(rootdir)
	#shutil.move('D:\ps\catchwikihow.py','D:\ps\\ready') 
	for i in range(0,len(list)):
		path = os.path.join(rootdir,list[i])
		if os.path.isfile(path) and path.endswith('html'):
			try:
				parseFile(path)
				shutil.move(path,'D:\ps\\ready')
			except :
				traceback.print_exc()
				print('解析',path,'失败了')
			#print(path)
			#break
